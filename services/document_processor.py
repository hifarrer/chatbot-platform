import os
import json
import PyPDF2
import docx
import requests
import re
import csv
from io import StringIO
from pathlib import Path
from openpyxl import load_workbook
from urllib.parse import urljoin, urlparse
from bs4 import BeautifulSoup
import httpx
import trafilatura
from collections import deque

class DocumentProcessor:
    def __init__(self):
        pass
    
    def process_document(self, file_path):
        """
        Process a document and extract text based on file type
        """
        print(f" DEBUG: Processing document: {file_path}")
        
        file_extension = Path(file_path).suffix.lower()
        print(f" DEBUG: File extension: {file_extension}")
        
        if file_extension == '.pdf':
            text = self._process_pdf(file_path)
        elif file_extension == '.docx':
            text = self._process_docx(file_path)
        elif file_extension == '.txt':
            text = self._process_txt(file_path)
        elif file_extension == '.json':
            text = self._process_json(file_path)
        elif file_extension == '.xlsx':
            text = self._process_xlsx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        print(f" DEBUG: Extracted {len(text)} characters from document")
        print(f" DEBUG: First 200 characters: {text[:200]}...")
        
        return text
    
    def _process_pdf(self, file_path):
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                print(f" DEBUG: PDF has {len(pdf_reader.pages)} pages")
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    text += page_text + "\n"
                    print(f"   Page {page_num + 1}: {len(page_text)} characters")
                    
        except Exception as e:
            raise Exception(f"Error processing PDF: {str(e)}")
        
        return text.strip()
    
    def _process_docx(self, file_path):
        """Extract text from DOCX file"""
        text = ""
        try:
            doc = docx.Document(file_path)
            print(f" DEBUG: DOCX has {len(doc.paragraphs)} paragraphs")
            
            for i, paragraph in enumerate(doc.paragraphs):
                paragraph_text = paragraph.text
                text += paragraph_text + "\n"
                if i < 5:  # Show first 5 paragraphs
                    print(f"   Paragraph {i + 1}: {paragraph_text[:100]}...")
                    
        except Exception as e:
            raise Exception(f"Error processing DOCX: {str(e)}")
        
        return text.strip()
    
    def _process_txt(self, file_path):
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
                print(f" DEBUG: TXT file read successfully")
        except UnicodeDecodeError:
            # Try with different encoding if UTF-8 fails
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
                    print(f" DEBUG: TXT file read with latin-1 encoding")
            except Exception as e:
                raise Exception(f"Error processing TXT file: {str(e)}")
        except Exception as e:
            raise Exception(f"Error processing TXT file: {str(e)}")
        
        return text.strip()
    
    def _process_json(self, file_path):
        """Extract text from JSON file - intelligently flattens JSON structure into readable text"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
                print(f" DEBUG: JSON file loaded successfully")
            
            # Convert JSON to readable text format
            text = self._json_to_text(data)
            print(f" DEBUG: Converted JSON to text: {len(text)} characters")
            
            return text.strip()
            
        except json.JSONDecodeError as e:
            raise Exception(f"Error processing JSON file: Invalid JSON format - {str(e)}")
        except Exception as e:
            raise Exception(f"Error processing JSON file: {str(e)}")
    
    def _json_to_text(self, data, indent=0, parent_key=''):
        """
        Recursively convert JSON data to readable text format
        Handles objects, arrays, and nested structures intelligently
        """
        text_parts = []
        indent_str = "  " * indent
        
        if isinstance(data, dict):
            # Handle dictionary/object
            for key, value in data.items():
                full_key = f"{parent_key}.{key}" if parent_key else key
                
                if isinstance(value, (dict, list)):
                    # Nested structure
                    text_parts.append(f"{indent_str}{key}:")
                    text_parts.append(self._json_to_text(value, indent + 1, full_key))
                else:
                    # Simple key-value pair
                    text_parts.append(f"{indent_str}{key}: {value}")
        
        elif isinstance(data, list):
            # Handle array
            for i, item in enumerate(data):
                if isinstance(item, (dict, list)):
                    # Nested structure in array
                    text_parts.append(f"{indent_str}Item {i + 1}:")
                    text_parts.append(self._json_to_text(item, indent + 1, f"{parent_key}[{i}]"))
                else:
                    # Simple array item
                    text_parts.append(f"{indent_str}- {item}")
        
        else:
            # Simple value (string, number, boolean, null)
            text_parts.append(f"{indent_str}{data}")
        
        return "\n".join(text_parts)
    
    def _process_xlsx(self, file_path):
        """
        Extract data from Excel file and convert to structured format
        Similar to Google Sheets processing - converts spreadsheet data to structured JSON/text
        """
        try:
            workbook = load_workbook(file_path, read_only=True, data_only=True)
            print(f" DEBUG: Excel file loaded successfully")
            print(f" DEBUG: Found {len(workbook.sheetnames)} sheet(s): {', '.join(workbook.sheetnames)}")
            
            all_sheets_text = []
            
            # Process each sheet in the workbook
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                print(f" DEBUG: Processing sheet: {sheet_name}")
                
                # Get all rows as lists
                rows = []
                for row in sheet.iter_rows(values_only=True):
                    # Filter out completely empty rows
                    if any(cell is not None and str(cell).strip() for cell in row):
                        # Convert all cells to strings, handling None values
                        cleaned_row = [str(cell).strip() if cell is not None else '' for cell in row]
                        rows.append(cleaned_row)
                
                if not rows:
                    print(f"   Sheet '{sheet_name}' is empty, skipping")
                    continue
                
                print(f"   Found {len(rows)} non-empty rows")
                
                # Assume first row is header
                if len(rows) > 0:
                    headers = rows[0]
                    data_rows = rows[1:]
                    
                    print(f"   Headers: {headers}")
                    print(f"   Data rows: {len(data_rows)}")
                    
                    # Convert to structured format similar to Google Sheets
                    structured_data = self._excel_to_structured_json(headers, data_rows, sheet_name)
                    
                    # Format for training
                    sheet_text = self._format_sheet_for_training(structured_data)
                    all_sheets_text.append(sheet_text)
            
            workbook.close()
            
            if not all_sheets_text:
                raise Exception("No data found in Excel file")
            
            # Combine all sheets
            final_text = "\n\n=== EXCEL WORKBOOK ===\n\n".join(all_sheets_text)
            print(f" DEBUG: Processed Excel file: {len(final_text)} characters total")
            
            return final_text.strip()
            
        except Exception as e:
            raise Exception(f"Error processing Excel file: {str(e)}")
    
    def _excel_to_structured_json(self, headers, data_rows, sheet_name):
        """
        Convert Excel data to structured JSON format optimized for AI training
        Similar to csv_to_structured_json but works with already-parsed data
        """
        print(f" DEBUG: Converting Excel sheet '{sheet_name}' to structured JSON")
        
        # Clean headers - remove empty columns
        cleaned_headers = []
        valid_column_indices = []
        for idx, header in enumerate(headers):
            if header and header.strip():
                cleaned_headers.append(header.strip())
                valid_column_indices.append(idx)
        
        if not cleaned_headers:
            return {
                "sheet_name": sheet_name,
                "description": "Empty sheet - no headers found",
                "data": []
            }
        
        print(f" DEBUG: Found {len(cleaned_headers)} valid columns: {cleaned_headers}")
        
        # Convert rows to dictionaries
        structured_rows = []
        for row in data_rows:
            row_dict = {}
            for header_idx, col_idx in enumerate(valid_column_indices):
                if col_idx < len(row):
                    value = row[col_idx]
                    # Only add non-empty values
                    if value and value.strip():
                        row_dict[cleaned_headers[header_idx]] = value.strip()
            
            # Only add rows that have at least one value
            if row_dict:
                structured_rows.append(row_dict)
        
        print(f" DEBUG: Converted {len(structured_rows)} non-empty rows to JSON")
        
        # Create structured JSON
        structured_data = {
            "sheet_name": sheet_name,
            "columns": cleaned_headers,
            "row_count": len(structured_rows),
            "data": structured_rows
        }
        
        return structured_data
    
    def chunk_text(self, text, chunk_size=1000, overlap=100):
        """
        Split text into chunks for better processing
        """
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), chunk_size - overlap):
            chunk = ' '.join(words[i:i + chunk_size])
            chunks.append(chunk)
            
            if i + chunk_size >= len(words):
                break
        
        print(f" DEBUG: Split text into {len(chunks)} chunks")
        return chunks
    
    def extract_google_doc_id(self, url):
        """
        Extract the document ID from a Google Docs URL
        Supports various Google Docs URL formats:
        - https://docs.google.com/document/d/{DOC_ID}/edit
        - https://docs.google.com/document/d/{DOC_ID}/
        - https://docs.google.com/document/d/{DOC_ID}
        """
        # Pattern to match Google Docs URLs
        patterns = [
            r'docs\.google\.com/document/d/([a-zA-Z0-9-_]+)',
            r'docs\.google\.com/document/u/\d+/d/([a-zA-Z0-9-_]+)',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, url)
            if match:
                doc_id = match.group(1)
                print(f" DEBUG: Extracted Google Doc ID: {doc_id}")
                return doc_id
        
        raise ValueError("Invalid Google Docs URL. Please provide a valid Google Docs link.")
    
    def fetch_google_doc(self, url):
        """
        Fetch text content from a Google Docs URL
        The document must be publicly accessible or shared with 'Anyone with the link can view'
        """
        print(f" DEBUG: Fetching Google Doc from: {url}")
        
        try:
            # Extract document ID from URL
            doc_id = self.extract_google_doc_id(url)
            
            # Construct the export URL for plain text format
            export_url = f"https://docs.google.com/document/d/{doc_id}/export?format=txt"
            print(f" DEBUG: Export URL: {export_url}")
            
            # Fetch the document
            response = requests.get(export_url, timeout=30)
            
            if response.status_code == 200:
                text = response.text
                print(f" DEBUG: Successfully fetched Google Doc: {len(text)} characters")
                print(f" DEBUG: First 200 characters: {text[:200]}...")
                return text.strip()
            elif response.status_code == 404:
                raise Exception("Google Doc not found. Please ensure the document exists and is shared with 'Anyone with the link can view'.")
            elif response.status_code == 403:
                raise Exception("Access denied. Please ensure the Google Doc is shared with 'Anyone with the link can view'.")
            else:
                raise Exception(f"Failed to fetch Google Doc. HTTP Status: {response.status_code}")
                
        except requests.RequestException as e:
            raise Exception(f"Network error while fetching Google Doc: {str(e)}")
        except ValueError as e:
            raise Exception(str(e))
        except Exception as e:
            raise Exception(f"Error fetching Google Doc: {str(e)}")
    
    def extract_google_sheet_id(self, url):
        """
        Extract the spreadsheet ID from a Google Sheets URL
        Supports various Google Sheets URL formats:
        - https://docs.google.com/spreadsheets/d/{SHEET_ID}/edit
        - https://docs.google.com/spreadsheets/d/{SHEET_ID}/
        - https://docs.google.com/spreadsheets/d/{SHEET_ID}
        """
        # Pattern to match Google Sheets URLs
        patterns = [
            r'docs\.google\.com/spreadsheets/d/([a-zA-Z0-9-_]+)',
            r'docs\.google\.com/spreadsheets/u/\d+/d/([a-zA-Z0-9-_]+)',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, url)
            if match:
                sheet_id = match.group(1)
                print(f" DEBUG: Extracted Google Sheet ID: {sheet_id}")
                return sheet_id
        
        raise ValueError("Invalid Google Sheets URL. Please provide a valid Google Sheets link.")
    
    def csv_to_structured_json(self, csv_text, sheet_name="Sheet"):
        """
        Convert CSV data to structured JSON format optimized for AI training
        Creates a format that's easy for OpenAI to understand and use
        """
        print(f" DEBUG: Converting CSV to structured JSON for sheet: {sheet_name}")
        
        # Parse CSV
        csv_reader = csv.DictReader(StringIO(csv_text))
        rows = list(csv_reader)
        
        if not rows:
            return {
                "sheet_name": sheet_name,
                "description": "Empty sheet",
                "data": []
            }
        
        # Get column headers
        headers = list(rows[0].keys())
        
        print(f" DEBUG: Found {len(rows)} rows with {len(headers)} columns")
        print(f" DEBUG: Headers: {headers}")
        
        # Create structured JSON
        structured_data = {
            "sheet_name": sheet_name,
            "columns": headers,
            "row_count": len(rows),
            "data": []
        }
        
        # Add all rows
        for idx, row in enumerate(rows):
            # Clean up empty values
            cleaned_row = {k: v for k, v in row.items() if v and v.strip()}
            if cleaned_row:  # Only add non-empty rows
                structured_data["data"].append(cleaned_row)
        
        print(f" DEBUG: Converted {len(structured_data['data'])} non-empty rows to JSON")
        
        return structured_data
    
    def fetch_google_sheet(self, url):
        """
        Fetch data from a Google Sheets URL and convert to structured JSON
        The spreadsheet must be publicly accessible or shared with 'Anyone with the link can view'
        
        Returns a formatted text representation that includes:
        - Sheet structure description
        - Data in both JSON format and human-readable text
        """
        print(f" DEBUG: Fetching Google Sheet from: {url}")
        
        try:
            # Extract spreadsheet ID from URL
            sheet_id = self.extract_google_sheet_id(url)
            
            # Construct the export URL for CSV format (default first sheet)
            # Note: gid=0 gets the first sheet, we could extend this to get all sheets
            export_url = f"https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=csv&gid=0"
            print(f" DEBUG: Export URL: {export_url}")
            
            # Fetch the spreadsheet
            response = requests.get(export_url, timeout=30)
            
            if response.status_code == 200:
                csv_text = response.text
                print(f" DEBUG: Successfully fetched Google Sheet: {len(csv_text)} characters")
                
                # Convert CSV to structured JSON
                structured_data = self.csv_to_structured_json(csv_text, "Main Sheet")
                
                # Create a comprehensive text representation for training
                training_text = self._format_sheet_for_training(structured_data)
                
                print(f" DEBUG: Generated training text: {len(training_text)} characters")
                return training_text
                
            elif response.status_code == 404:
                raise Exception("Google Sheet not found. Please ensure the spreadsheet exists and is shared with 'Anyone with the link can view'.")
            elif response.status_code == 403:
                raise Exception("Access denied. Please ensure the Google Sheet is shared with 'Anyone with the link can view'.")
            else:
                raise Exception(f"Failed to fetch Google Sheet. HTTP Status: {response.status_code}")
                
        except requests.RequestException as e:
            raise Exception(f"Network error while fetching Google Sheet: {str(e)}")
        except ValueError as e:
            raise Exception(str(e))
        except Exception as e:
            raise Exception(f"Error fetching Google Sheet: {str(e)}")
    
    def _format_sheet_for_training(self, structured_data):
        """
        Format structured sheet data into a comprehensive text format for AI training
        Combines JSON structure with human-readable descriptions
        """
        lines = []
        
        # Add header with sheet information
        lines.append(f"=== Google Sheet: {structured_data['sheet_name']} ===")
        lines.append(f"Total Rows: {structured_data['row_count']}")
        lines.append(f"Columns: {', '.join(structured_data['columns'])}")
        lines.append("")
        
        # Add JSON representation (for structured understanding)
        lines.append("=== Structured Data (JSON Format) ===")
        lines.append(json.dumps(structured_data, indent=2, ensure_ascii=False))
        lines.append("")
        
        # Add human-readable table format
        lines.append("=== Human-Readable Format ===")
        
        if structured_data['data']:
            # Create a readable list of items
            for idx, row in enumerate(structured_data['data'], 1):
                lines.append(f"\nEntry {idx}:")
                for key, value in row.items():
                    lines.append(f"  {key}: {value}")
        
        lines.append("")
        lines.append("=== End of Sheet Data ===")
        
        return "\n".join(lines)
    
    def scrape_website(self, url, max_pages=50, timeout=120):
        """
        Scrape a website and extract clean text content from up to max_pages pages.
        
        Strategy:
        1. Try to parse sitemap.xml first for fastest discovery
        2. If no sitemap, crawl from homepage (breadth-first)
        3. Use httpx for fetching (fast, modern)
        4. Use trafilatura for content extraction (best-in-class)
        5. Limit to max_pages to respect OpenAI context limits
        
        Args:
            url (str): The website URL to scrape
            max_pages (int): Maximum number of pages to scrape (default 50)
            timeout (int): Maximum time in seconds for entire operation (default 120)
            
        Returns:
            str: Combined text from all scraped pages
        """
        print(f" DEBUG: Starting website scraping for: {url}")
        print(f" DEBUG: Max pages: {max_pages}, Timeout: {timeout}s")
        
        import time
        start_time = time.time()
        
        try:
            # Normalize URL
            if not url.startswith(('http://', 'https://')):
                url = 'https://' + url
            
            parsed_url = urlparse(url)
            base_domain = f"{parsed_url.scheme}://{parsed_url.netloc}"
            
            print(f" DEBUG: Base domain: {base_domain}")
            
            # Step 1: Try to get URLs from sitemap.xml
            urls_to_scrape = self._get_urls_from_sitemap(base_domain, max_pages)
            
            # Step 2: If no sitemap or insufficient URLs, crawl the site
            if len(urls_to_scrape) < max_pages:
                print(f" DEBUG: Sitemap provided {len(urls_to_scrape)} URLs, crawling for more...")
                crawled_urls = self._crawl_website(url, base_domain, max_pages - len(urls_to_scrape))
                urls_to_scrape.extend(crawled_urls)
            
            # Remove duplicates and limit to max_pages
            urls_to_scrape = list(dict.fromkeys(urls_to_scrape))[:max_pages]
            
            print(f" DEBUG: Will scrape {len(urls_to_scrape)} pages")
            
            # Step 3: Fetch and extract content from each URL
            all_content = []
            successful_scrapes = 0
            
            with httpx.Client(timeout=30.0, follow_redirects=True) as client:
                for idx, page_url in enumerate(urls_to_scrape, 1):
                    # Check timeout
                    if time.time() - start_time > timeout:
                        print(f" WARNING: Timeout reached after {idx-1} pages")
                        break
                    
                    try:
                        print(f" DEBUG: [{idx}/{len(urls_to_scrape)}] Fetching: {page_url}")
                        
                        # Fetch the page
                        response = client.get(
                            page_url,
                            headers={
                                'User-Agent': 'Mozilla/5.0 (compatible; OwlbeeChatbotTrainer/1.0; +https://owlbee.ai)'
                            }
                        )
                        
                        if response.status_code == 200:
                            # Extract clean text using trafilatura
                            extracted_text = trafilatura.extract(
                                response.text,
                                include_comments=False,
                                include_tables=True,
                                no_fallback=False
                            )
                            
                            if extracted_text and len(extracted_text.strip()) > 100:
                                # Add page metadata
                                page_content = f"\n\n=== PAGE: {page_url} ===\n\n{extracted_text}\n\n=== END OF PAGE ===\n"
                                all_content.append(page_content)
                                successful_scrapes += 1
                                print(f"   ✓ Extracted {len(extracted_text)} characters")
                            else:
                                print(f"   ✗ No content extracted or content too short")
                        else:
                            print(f"   ✗ HTTP {response.status_code}")
                        
                        # Be polite - rate limiting
                        time.sleep(0.5)
                        
                    except Exception as e:
                        print(f"   ✗ Error fetching {page_url}: {str(e)}")
                        continue
            
            if not all_content:
                raise Exception("No content could be extracted from the website. The site may be blocking automated access or the content may not be accessible.")
            
            # Combine all content
            combined_text = "\n".join(all_content)
            
            # Add summary header
            summary = f"""=== WEBSITE SCRAPE SUMMARY ===
Source: {url}
Base Domain: {base_domain}
Pages Scraped: {successful_scrapes} of {len(urls_to_scrape)} attempted
Total Characters: {len(combined_text)}
Scraped At: {time.strftime('%Y-%m-%d %H:%M:%S')}
=== END SUMMARY ===

"""
            
            final_text = summary + combined_text
            
            elapsed_time = time.time() - start_time
            print(f" DEBUG: Scraping complete!")
            print(f"   - Successful: {successful_scrapes}/{len(urls_to_scrape)} pages")
            print(f"   - Total content: {len(final_text)} characters")
            print(f"   - Time elapsed: {elapsed_time:.1f}s")
            
            return final_text
            
        except Exception as e:
            print(f" ERROR: Website scraping failed: {str(e)}")
            raise Exception(f"Failed to scrape website: {str(e)}")
    
    def _get_urls_from_sitemap(self, base_url, max_urls=50):
        """
        Try to extract URLs from sitemap.xml
        Returns list of URLs found in sitemap
        Uses multiple parsing approaches for maximum compatibility
        """
        urls = []
        sitemap_urls = [
            f"{base_url}/sitemap.xml",
            f"{base_url}/sitemap_index.xml",
            f"{base_url}/sitemap1.xml"
        ]
        
        print(f" DEBUG: Checking for sitemap...")
        
        try:
            with httpx.Client(timeout=10.0, follow_redirects=True) as client:
                for sitemap_url in sitemap_urls:
                    try:
                        response = client.get(sitemap_url)
                        if response.status_code == 200:
                            print(f"   ✓ Found sitemap: {sitemap_url}")
                            
                            # Try multiple parsing approaches
                            urls_found = self._parse_sitemap_content(response.content, max_urls)
                            if urls_found:
                                urls.extend(urls_found)
                                print(f"   ✓ Extracted {len(urls_found)} URLs from sitemap")
                                if len(urls) >= max_urls:
                                    return urls[:max_urls]
                    except Exception as e:
                        print(f"   ✗ Error parsing sitemap {sitemap_url}: {str(e)}")
                        continue
        except Exception as e:
            print(f"   ✗ No sitemap found: {str(e)}")
        
        return urls[:max_urls]
    
    def _parse_sitemap_content(self, content, max_urls=50):
        """
        Parse sitemap content using multiple approaches for maximum compatibility
        """
        urls = []
        
        # Approach 1: Try BeautifulSoup with built-in XML parser
        try:
            soup = BeautifulSoup(content, 'xml')
            loc_tags = soup.find_all('loc')
            for loc in loc_tags:
                url = loc.text.strip()
                if url and url.startswith('http'):
                    urls.append(url)
                    if len(urls) >= max_urls:
                        break
            if urls:
                return urls
        except Exception as e:
            print(f"   ✗ BeautifulSoup XML parsing failed: {str(e)}")
        
        # Approach 2: Try BeautifulSoup with HTML parser (fallback)
        try:
            soup = BeautifulSoup(content, 'html.parser')
            loc_tags = soup.find_all('loc')
            for loc in loc_tags:
                url = loc.text.strip()
                if url and url.startswith('http'):
                    urls.append(url)
                    if len(urls) >= max_urls:
                        break
            if urls:
                return urls
        except Exception as e:
            print(f"   ✗ BeautifulSoup HTML parsing failed: {str(e)}")
        
        # Approach 3: Simple regex extraction (most compatible)
        try:
            import re
            # Look for URLs in <loc> tags using regex
            url_pattern = r'<loc>(https?://[^<]+)</loc>'
            matches = re.findall(url_pattern, content.decode('utf-8', errors='ignore'))
            for url in matches:
                url = url.strip()
                if url and url.startswith('http'):
                    urls.append(url)
                    if len(urls) >= max_urls:
                        break
            if urls:
                return urls
        except Exception as e:
            print(f"   ✗ Regex parsing failed: {str(e)}")
        
        return []
    
    def _crawl_website(self, start_url, base_domain, max_urls=50):
        """
        Crawl website starting from start_url using breadth-first search.
        Only follows internal links within the same domain.
        
        Returns list of URLs discovered
        """
        print(f" DEBUG: Crawling website from: {start_url}")
        
        visited = set()
        to_visit = deque([start_url])
        discovered_urls = []
        
        try:
            with httpx.Client(timeout=10.0, follow_redirects=True) as client:
                while to_visit and len(discovered_urls) < max_urls:
                    current_url = to_visit.popleft()
                    
                    # Skip if already visited
                    if current_url in visited:
                        continue
                    
                    visited.add(current_url)
                    
                    try:
                        print(f"   Crawling: {current_url}")
                        response = client.get(
                            current_url,
                            headers={
                                'User-Agent': 'Mozilla/5.0 (compatible; OwlbeeChatbotTrainer/1.0; +https://owlbee.ai)'
                            }
                        )
                        
                        if response.status_code == 200 and 'text/html' in response.headers.get('content-type', ''):
                            discovered_urls.append(current_url)
                            
                            # Parse HTML to find more links
                            soup = BeautifulSoup(response.text, 'html.parser')
                            
                            # Find all links
                            for link in soup.find_all('a', href=True):
                                href = link['href']
                                
                                # Convert relative URLs to absolute
                                absolute_url = urljoin(current_url, href)
                                
                                # Only add if it's from the same domain and not visited
                                parsed = urlparse(absolute_url)
                                if (parsed.netloc == urlparse(base_domain).netloc and
                                    absolute_url not in visited and
                                    absolute_url not in to_visit and
                                    not self._should_skip_url(absolute_url)):
                                    to_visit.append(absolute_url)
                        
                        # Be polite
                        time.sleep(0.3)
                        
                    except Exception as e:
                        print(f"   ✗ Error crawling {current_url}: {str(e)}")
                        continue
        except Exception as e:
            print(f" ERROR: Crawling failed: {str(e)}")
        
        print(f"   ✓ Discovered {len(discovered_urls)} URLs via crawling")
        return discovered_urls
    
    def _should_skip_url(self, url):
        """
        Check if URL should be skipped (e.g., files, common non-content pages)
        """
        skip_extensions = [
            '.pdf', '.jpg', '.jpeg', '.png', '.gif', '.zip', '.tar', '.gz',
            '.mp4', '.mp3', '.avi', '.mov', '.doc', '.docx', '.xls', '.xlsx',
            '.css', '.js', '.xml', '.json', '.ico', '.svg', '.woff', '.ttf'
        ]
        
        skip_patterns = [
            '/feed/', '/rss/', '/atom/', '/wp-json/', '/api/',
            '/login', '/logout', '/signin', '/signup', '/register',
            '/cart', '/checkout', '/account', '/admin',
            '#', 'javascript:', 'mailto:', 'tel:'
        ]
        
        url_lower = url.lower()
        
        # Skip if has file extension we want to avoid
        for ext in skip_extensions:
            if url_lower.endswith(ext):
                return True
        
        # Skip if matches patterns
        for pattern in skip_patterns:
            if pattern in url_lower:
                return True
        
        return False 